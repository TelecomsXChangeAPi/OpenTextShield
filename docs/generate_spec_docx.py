#!/usr/bin/env python3
"""Generate OpenTextShield Hardware Spec Sheet as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, header_color="1B4F72", stripe=True):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if stripe and r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    # Set cell padding
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    return table

def main():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ===== COVER SECTION =====
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OpenTextShield")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Hardware Specification Sheet")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph("")

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("AI-Powered SMS Classification Engine")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Cover details
    details = [
        ("Product", "OpenTextShield mBERT v2.5"),
        ("Target Throughput", "25 Transactions Per Second (TPS)"),
        ("Daily Capacity", "2,160,000 lookups/day"),
        ("Max Response Latency", "< 500 ms"),
        ("Platform", "VMware vSphere / ESXi"),
        ("Date", "March 3, 2026"),
        ("Prepared by", "TelecomsXChange (TCXC)"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
        run = p.add_run(value)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS (manual) =====
    toc_title = doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Architecture Overview",
        "1. Deployment Target",
        "2. Model Specifications",
        "3. Hardware Configurations",
        "    3A. Single VM with GPU (Recommended)",
        "    3B. Single VM, CPU-Only",
        "    3C. Two VMs with Load Balancer (High Availability)",
        "4. Configuration Summary",
        "5. Software Stack",
        "6. Network & Firewall Requirements",
        "7. Capacity Planning",
        "8. Deployment Steps",
        "9. Monitoring & Health Checks",
        "10. Support & Contact",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if item.startswith("    "):
            p.paragraph_format.left_indent = Cm(1.5)

    doc.add_page_break()

    # ===== ARCHITECTURE OVERVIEW =====
    doc.add_heading("Architecture Overview", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "OpenTextShield uses an async inference pipeline that offloads model inference to a dedicated "
        "thread pool (ThreadPoolExecutor). This keeps the FastAPI event loop responsive and enables "
        "concurrent request processing within a single worker process."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Per-request processing pipeline:")
    run.bold = True

    add_styled_table(doc,
        ["Stage", "Operation", "Time"],
        [
            ["1", "Enhanced preprocessing (Unicode, homoglyph, URL features)", "~2-5 ms"],
            ["2", "BERT tokenization (encode_plus, max_length=512)", "~1-3 ms"],
            ["3", "Model inference (12-layer BERT, torch.no_grad)", "~40-150 ms"],
            ["4", "Post-processing (softmax, label mapping)", "< 1 ms"],
            ["5", "Audit logging (fire-and-forget, non-blocking)", "0 ms*"],
        ],
        header_color="1A5276"
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("*")
    run.italic = True
    run.font.size = Pt(8)
    p.add_run(" Audit logging runs asynchronously after the response is sent.").font.size = Pt(8)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Concurrency model:")
    run.bold = True

    add_styled_table(doc,
        ["Device", "GIL Released?", "Concurrent Inferences/Worker", "Effective TPS/Worker"],
        [
            ["NVIDIA CUDA (T4/L4)", "Yes", "Up to 4 (thread pool)", "~30-50"],
            ["Apple MPS (benchmarked)", "Partial", "Up to 4 (thread pool)", "~17"],
            ["CPU", "No (GIL held)", "1 effective", "~7-10"],
        ],
        header_color="922B21"
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Benchmark reference: ")
    run.bold = True
    p.add_run(
        "On Apple MPS (M-series), a single worker achieved 17.4 TPS with 10 concurrent requests "
        "(1.8x improvement over sequential). CUDA GPUs (T4/L4) release the GIL more completely "
        "during kernel execution, so the concurrency multiplier is expected to be higher (~2-3x). "
        "Health checks remained responsive at 6.2ms average during concurrent load. "
        "On CPU, the GIL limits parallelism so multiple worker processes are needed instead."
    )

    doc.add_page_break()

    # ===== 1. DEPLOYMENT TARGET =====
    doc.add_heading("1. Deployment Target", level=1)

    add_styled_table(doc,
        ["Parameter", "Requirement"],
        [
            ["Throughput", "25 transactions per second (TPS) sustained"],
            ["Maximum Response Latency", "< 500 ms per request"],
            ["Daily Lookup Capacity", "2,160,000 lookups/day (25 TPS x 86,400 sec)"],
            ["Availability", "24/7 continuous operation"],
            ["Virtualization Platform", "VMware vSphere / ESXi"],
        ]
    )

    doc.add_paragraph("")

    # ===== 2. MODEL SPECIFICATIONS =====
    doc.add_heading("2. Model Specifications", level=1)

    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Model Architecture", "BERT-base-multilingual-cased (mBERT)"],
            ["Total Parameters", "~110 million"],
            ["Model File Size", "679 MB (.pth format)"],
            ["Runtime Memory per Instance", "~2 GB"],
            ["Supported Languages", "104+ languages"],
            ["Classification Labels", "ham, spam, phishing"],
            ["Max Input Length", "512 tokens"],
            ["Framework", "PyTorch 2.7.1"],
            ["GPU Acceleration", "NVIDIA CUDA (recommended), CPU fallback supported"],
        ]
    )

    doc.add_paragraph("")

    # ===== 3. HARDWARE CONFIGURATIONS =====
    doc.add_heading("3. Hardware Configurations", level=1)

    # --- Option A ---
    doc.add_heading("Option A - Single VM with GPU (Recommended)", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Best for: ")
    run.bold = True
    p.add_run("Lowest latency, simplest architecture, easiest to maintain.")

    doc.add_paragraph("")

    add_styled_table(doc,
        ["Resource", "Specification"],
        [
            ["vCPUs", "4 vCPUs (Intel Xeon Scalable 3rd Gen+ or AMD EPYC 7003+)"],
            ["RAM", "8 GB DDR4 ECC"],
            ["GPU", "1x NVIDIA T4 (16 GB VRAM) or NVIDIA L4 (24 GB VRAM)"],
            ["Storage", "40 GB SSD (NVMe preferred)"],
            ["Network", "1 Gbps virtual NIC (vmxnet3)"],
            ["OS", "Ubuntu 24.04 LTS (Server)"],
            ["VMware Requirement", "GPU Passthrough (DirectPath I/O) or NVIDIA vGPU"],
        ]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Performance Profile:")
    run.bold = True
    run.font.size = Pt(10)

    add_styled_table(doc,
        ["Metric", "Expected Value"],
        [
            ["Per-request latency (end-to-end)", "45 - 85 ms"],
            ["Sustained throughput", "30 - 50 TPS"],
            ["Headroom above 25 TPS target", "~20 - 100%"],
            ["Uvicorn workers", "1 (use 2 for additional headroom)"],
            ["Inference threads per worker", "4 (ThreadPoolExecutor)"],
            ["Model memory in VRAM", "~2 GB of 16 GB available"],
        ],
        header_color="1A5276"
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Why 1 worker is sufficient: ")
    run.bold = True
    p.add_run(
        "CUDA releases Python's GIL during GPU kernel execution. The 4-thread inference pool "
        "allows concurrent forward passes to overlap on the GPU. Benchmarked at 17.4 TPS on "
        "Apple MPS (1.8x over sequential); CUDA T4/L4 expected to achieve 30-50 TPS due to "
        "more complete GIL release. For production headroom, use --workers 2."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("VMware GPU Notes:")
    run.bold = True
    notes_a = [
        "Requires ESXi 7.0+ with GPU passthrough enabled",
        "Compatible GPUs: NVIDIA T4, L4, A2, A10, A30",
        "NVIDIA vGPU licensing is an alternative to full passthrough",
        "Install NVIDIA CUDA drivers inside guest VM",
    ]
    for note in notes_a:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_paragraph("")

    # --- Option B ---
    doc.add_heading("Option B - Single VM, CPU-Only (No GPU Required)", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Best for: ")
    run.bold = True
    p.add_run("Environments without GPU passthrough capability.")

    doc.add_paragraph("")

    add_styled_table(doc,
        ["Resource", "Specification"],
        [
            ["vCPUs", "16 vCPUs (Intel Xeon Scalable 3rd Gen+ or AMD EPYC 7003+)"],
            ["RAM", "24 GB DDR4 ECC"],
            ["GPU", "None required"],
            ["Storage", "40 GB SSD (NVMe preferred)"],
            ["Network", "1 Gbps virtual NIC (vmxnet3)"],
            ["OS", "Ubuntu 24.04 LTS (Server)"],
        ]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Performance Profile:")
    run.bold = True

    add_styled_table(doc,
        ["Metric", "Expected Value"],
        [
            ["Per-request latency (end-to-end)", "105 - 260 ms"],
            ["Sustained throughput", "28 - 40 TPS"],
            ["Headroom above target", "~12 - 60%"],
            ["Uvicorn workers", "4"],
            ["Model instances in RAM", "4 (4 x 2 GB = 8 GB model memory)"],
        ],
        header_color="1A5276"
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Why 4 workers: ")
    run.bold = True
    p.add_run(
        "On CPU, Python's GIL prevents true thread-level parallelism for compute-bound inference. "
        "Each uvicorn worker is a separate process with its own GIL, so 4 workers = 4 truly parallel "
        "inferences. The async thread pool still keeps each worker's event loop responsive for health checks."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("CPU Guidance:")
    run.bold = True
    notes_b = [
        "Higher single-thread clock speed improves per-request latency",
        "Minimum recommended: 2.4 GHz base clock",
        "AVX-512 instruction support beneficial for PyTorch inference",
        "Hyper-threading should be enabled",
    ]
    for note in notes_b:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_paragraph("")

    # --- Option C ---
    doc.add_heading("Option C - Two VMs with Load Balancer (High Availability)", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Best for: ")
    run.bold = True
    p.add_run("Production environments requiring redundancy and failover.")

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Application VMs (x2):")
    run.bold = True
    run.font.size = Pt(10)

    add_styled_table(doc,
        ["Resource", "Per VM"],
        [
            ["vCPUs", "8 vCPUs"],
            ["RAM", "12 GB DDR4 ECC"],
            ["GPU", "None required"],
            ["Storage", "30 GB SSD"],
            ["Network", "1 Gbps virtual NIC (vmxnet3)"],
            ["OS", "Ubuntu 24.04 LTS (Server)"],
        ]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Load Balancer VM (x1):")
    run.bold = True
    run.font.size = Pt(10)

    add_styled_table(doc,
        ["Resource", "Specification"],
        [
            ["vCPUs", "2 vCPUs"],
            ["RAM", "2 GB DDR4"],
            ["Storage", "10 GB SSD"],
            ["Network", "1 Gbps virtual NIC (vmxnet3)"],
            ["Software", "Nginx (included with OpenTextShield)"],
            ["Algorithm", "Least Connections"],
        ]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Performance Profile:")
    run.bold = True

    add_styled_table(doc,
        ["Metric", "Expected Value"],
        [
            ["Per-request latency (end-to-end)", "105 - 210 ms"],
            ["Sustained throughput", "28 - 40 TPS"],
            ["Headroom above target", "~12 - 60%"],
            ["Uvicorn workers per VM", "2"],
            ["Total model instances", "4 (2 per VM)"],
            ["Fault tolerance", "1 VM can fail; remaining handles ~14-20 TPS"],
        ],
        header_color="1A5276"
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Total Physical Resources (Option C):")
    run.bold = True

    add_styled_table(doc,
        ["Resource", "Total"],
        [
            ["vCPUs", "18"],
            ["RAM", "26 GB"],
            ["Storage", "70 GB"],
            ["VMs", "3"],
        ]
    )

    doc.add_page_break()

    # ===== 4. CONFIGURATION SUMMARY =====
    doc.add_heading("4. Configuration Summary", level=1)

    add_styled_table(doc,
        ["Option", "vCPUs", "RAM", "GPU", "Latency", "TPS Capacity", "Redundancy", "Complexity"],
        [
            ["A: GPU", "4", "8 GB", "1x T4/L4", "45-85 ms", "30-50", "None", "Low"],
            ["B: CPU", "16", "24 GB", "None", "105-260 ms", "28-40", "None", "Low"],
            ["C: HA", "18", "26 GB", "None", "105-210 ms", "28-40", "Yes", "Medium"],
        ]
    )

    doc.add_paragraph("")

    # ===== 5. SOFTWARE STACK =====
    doc.add_heading("5. Software Stack", level=1)

    add_styled_table(doc,
        ["Component", "Version", "Purpose"],
        [
            ["Ubuntu Server", "24.04 LTS", "Operating system"],
            ["Python", "3.12.x", "Runtime"],
            ["PyTorch", "2.7.1", "Model inference engine"],
            ["Transformers", "4.53.0+", "Tokenizer and model loading"],
            ["FastAPI", "Latest", "REST API framework"],
            ["Uvicorn", "Latest", "ASGI application server"],
            ["Nginx", "Latest", "Load balancer (Option C only)"],
            ["Docker", "24.x+ (optional)", "Container deployment"],
            ["NVIDIA CUDA", "12.x (Option A only)", "GPU acceleration"],
        ]
    )

    doc.add_paragraph("")

    # ===== 6. NETWORK & FIREWALL =====
    doc.add_heading("6. Network & Firewall Requirements", level=1)

    add_styled_table(doc,
        ["Port", "Protocol", "Direction", "Purpose"],
        [
            ["8002", "TCP", "Inbound", "API endpoint (HTTP REST)"],
            ["8080", "TCP", "Inbound", "Web frontend (optional)"],
            ["443", "TCP", "Outbound", "Package downloads during setup"],
            ["80", "TCP", "Outbound", "Package downloads during setup"],
        ]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("API Endpoints:")
    run.bold = True

    endpoints = [
        "API Endpoint: POST /predict/",
        "Health Check: GET /health",
        "API Documentation: GET /docs (Swagger UI)",
    ]
    for ep in endpoints:
        doc.add_paragraph(ep, style='List Bullet')

    doc.add_paragraph("")

    # ===== 7. CAPACITY PLANNING =====
    doc.add_heading("7. Capacity Planning", level=1)

    doc.add_heading("Daily Volume at Various TPS Levels", level=2)

    add_styled_table(doc,
        ["TPS", "Hourly", "Daily", "Monthly (30d)"],
        [
            ["10", "36,000", "864,000", "25,920,000"],
            ["25 (target)", "90,000", "2,160,000", "64,800,000"],
            ["50", "180,000", "4,320,000", "129,600,000"],
            ["100", "360,000", "8,640,000", "259,200,000"],
        ]
    )

    doc.add_paragraph("")

    doc.add_heading("GPU Scaling (per VM with NVIDIA T4/L4)", level=2)

    add_styled_table(doc,
        ["VMs", "Workers", "Approx. Max TPS", "Daily Capacity"],
        [
            ["1", "1", "30 - 50", "2.6M - 4.3M"],
            ["2", "2", "60 - 100", "5.2M - 8.6M"],
            ["4", "4", "120 - 200", "10.4M - 17.3M"],
        ]
    )

    doc.add_paragraph("")

    doc.add_heading("CPU-Only Scaling (per VM added, 4 workers each)", level=2)

    add_styled_table(doc,
        ["VMs", "Workers", "Approx. Max TPS", "Daily Capacity"],
        [
            ["1", "4", "28 - 40", "2.4M - 3.5M"],
            ["2", "8", "56 - 80", "4.8M - 6.9M"],
            ["4", "16", "112 - 160", "9.7M - 13.8M"],
            ["10", "40", "280 - 400", "24.2M - 34.6M"],
        ]
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Throughput scales approximately linearly with the number of VM instances.").italic = True

    doc.add_page_break()

    # ===== 8. DEPLOYMENT STEPS =====
    doc.add_heading("8. Deployment Steps (Quick Start)", level=1)

    steps = [
        ("Step 1 - Install system dependencies",
         "sudo apt update && sudo apt install -y python3.12 python3.12-venv git curl"),
        ("Step 2 - Clone repository",
         "git clone https://github.com/TelecomsXChangeAPi/OpenTextShield.git\ncd OpenTextShield"),
        ("Step 3 - Create virtual environment",
         "python3.12 -m venv ots\nsource ots/bin/activate\npip install --upgrade pip\npip install -r requirements.txt"),
        ("Step 4 - Start the server",
         "# Option A (GPU): 1 worker (thread pool handles concurrency)\n# Option B (CPU): 4 workers (separate processes for parallelism)\nuvicorn src.api_interface.main:app --host 0.0.0.0 --port 8002 --workers 4"),
        ("Step 5 - Verify deployment",
         'curl -X POST "http://localhost:8002/predict/" \\\n  -H "Content-Type: application/json" \\\n  -d \'{"text":"Congratulations! You won a free prize!","model":"ots-mbert"}\''),
    ]

    for title, cmd in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)

        code_p = doc.add_paragraph()
        run = code_p.add_run(cmd)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        code_p.paragraph_format.left_indent = Cm(1)
        doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Docker Alternative:")
    run.bold = True
    code_p = doc.add_paragraph()
    run = code_p.add_run("docker-compose up -d")
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    code_p.paragraph_format.left_indent = Cm(1)

    doc.add_paragraph("")

    # ===== 9. MONITORING =====
    doc.add_heading("9. Monitoring & Health Checks", level=1)

    add_styled_table(doc,
        ["Endpoint", "Method", "Expected Response"],
        [
            ["/health", "GET", "200 OK with model status"],
            ["/docs", "GET", "Swagger API documentation"],
            ["/predict/", "POST", "Classification result with timing"],
        ]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Recommended Monitoring Alerts:")
    run.bold = True

    alerts = [
        "Response time per request (alert if > 500 ms)",
        "Requests per second (alert if sustained > 80% of capacity)",
        "CPU and memory utilization per VM",
        "GPU utilization % (Option A - alert if sustained > 90%)",
        "Error rate (alert if > 1%)",
        "Model load status on application startup",
        "Thread pool queue depth (alert if inference requests are queuing)",
    ]
    for alert in alerts:
        doc.add_paragraph(alert, style='List Bullet')

    doc.add_paragraph("")

    # ===== 10. SUPPORT =====
    doc.add_heading("10. Support & Contact", level=1)

    add_styled_table(doc,
        ["Item", "Detail"],
        [
            ["Product", "OpenTextShield"],
            ["Vendor", "TelecomsXChange (TCXC)"],
            ["Repository", "https://github.com/TelecomsXChangeAPi/OpenTextShield"],
            ["API Docs", "http://<server-ip>:8002/docs"],
            ["Model Version", "OTS mBERT v2.5"],
        ]
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Disclaimer
    p = doc.add_paragraph()
    run = p.add_run(
        "This specification is based on performance benchmarks conducted with OpenTextShield v2.5 "
        "using the async inference pipeline. Actual performance may vary based on VMware host "
        "configuration, resource contention, network conditions, and workload characteristics. "
        "Benchmark validation on target hardware is recommended prior to production deployment."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "OpenTextShield_Hardware_Spec_Sheet_25TPS.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    main()
